from pathlib import Path
import sys

from lxml import etree
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(r"C:\Users\redma\Documents\Portfolio Arcane")
SKILL_ROOT = Path(
    r"C:\Users\redma\.codex\plugins\cache\openai-primary-runtime\documents"
    r"\26.715.12143\skills\documents"
)
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


OUT_DIR = ROOT / "output" / "documents"
WORK_DIR = ROOT / "tmp" / "docs"
RAW_OUT = WORK_DIR / "Malaska_Resume_2026_Google_Docs_raw.docx"
FINAL_OUT = OUT_DIR / "Malaska_Resume_2026_Google_Docs.docx"
OUT_DIR.mkdir(parents=True, exist_ok=True)
WORK_DIR.mkdir(parents=True, exist_ok=True)


# google_docs_default with deliberate branded-resume overrides based on the
# user's 2023 resume: gray page, navy/red palette, monospaced section labels,
# a precise two-column alignment grid, and compact one-page geometry.
PAGE_MARGIN = Inches(0.45)
PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = PAGE_WIDTH_DXA - 2 * int(PAGE_MARGIN.twips)
LABEL_WIDTH_DXA = 3600  # 2.50 in; matches the broad label rail in the reference
DETAIL_WIDTH_DXA = CONTENT_WIDTH_DXA - LABEL_WIDTH_DXA

FONT = "Arial"
LABEL_FONT = "Inconsolata"
INK = RGBColor(0x11, 0x19, 0x36)
BODY = RGBColor(0x34, 0x38, 0x5F)
ACCENT = RGBColor(0xFF, 0x31, 0x50)
MUTED = RGBColor(0xB9, 0xB9, 0xC1)
BACKGROUND = "D6D7DA"
BORDER = "BFC0C5"


def set_run_font(run, size, *, bold=False, italic=False, color=BODY, font_name=FONT):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)
    return run


def set_paragraph(
    paragraph,
    *,
    before=0,
    after=0,
    line=1.0,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    keep_with_next=False,
):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line
    fmt.alignment = alignment
    fmt.keep_with_next = keep_with_next
    fmt.widow_control = False
    return paragraph


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    return paragraph


def set_cell_border(cell, *, top=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "right", "bottom", "insideH", "insideV"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "nil")
    if top:
        element = tc_borders.find(qn("w:top"))
        if element is None:
            element = OxmlElement("w:top")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(top.get("sz", 4)))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), top.get("color", BORDER))
    else:
        element = tc_borders.find(qn("w:top"))
        if element is None:
            element = OxmlElement("w:top")
            tc_borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_margins(cell, *, top=100, bottom=100, start=0, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {
        "top": top,
        "bottom": bottom,
        "start": start,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill=BACKGROUND):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_document_background(doc, fill=BACKGROUND):
    document = doc._element
    background = document.find(qn("w:background"))
    if background is None:
        background = OxmlElement("w:background")
        document.insert(0, background)
    background.set(qn("w:color"), fill)
    settings = doc.settings._element
    display = settings.find(qn("w:displayBackgroundShape"))
    if display is None:
        display = OxmlElement("w:displayBackgroundShape")
        settings.append(display)


def add_printing_page_background(section, fill=BACKGROUND):
    """Add a header-anchored VML rectangle so the gray page prints/renders."""
    header = section.header
    paragraph = clear_paragraph(header.paragraphs[0])
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = OxmlElement("w:r")
    pict = OxmlElement("w:pict")
    rect = etree.Element(
        "{urn:schemas-microsoft-com:vml}rect",
        nsmap={"v": "urn:schemas-microsoft-com:vml"},
    )
    rect.set("id", "ResumePageBackground")
    rect.set(
        "style",
        "position:absolute;left:0;top:0;width:8.5in;height:11in;"
        "z-index:-251654144;mso-position-horizontal:left;"
        "mso-position-horizontal-relative:page;mso-position-vertical:top;"
        "mso-position-vertical-relative:page;mso-wrap-edited:f",
    )
    rect.set("fillcolor", f"#{fill}")
    rect.set("stroked", "f")
    pict.append(rect)
    run.append(pict)
    paragraph._p.append(run)


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_section_label(cell, text):
    p = clear_paragraph(cell.paragraphs[0])
    set_paragraph(p, line=1.0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_run_font(p.add_run(text.upper()), 15.0, color=INK, font_name=LABEL_FONT)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    # Encode the google_docs_default hierarchy, even though this resume uses
    # the compact named paragraph roles below.
    for style_name, size, color, before, after in (
        ("Heading 1", 20, RGBColor(0, 0, 0), 20, 6),
        ("Heading 2", 16, RGBColor(0, 0, 0), 18, 6),
        ("Heading 3", 14, RGBColor(0x43, 0x43, 0x43), 16, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = False
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_real_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "+")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "100")
    spacing.set(qn("w:line"), "276")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "FF3150")
    r_pr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    r_pr.append(sz)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.insert(0, num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_hyperlink(paragraph, text, url, *, size=10.5, color=BODY):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), FONT)
    r_pr.append(fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    r_pr.append(color_el)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_contact_line(cell, icon, value, url):
    p = cell.add_paragraph()
    set_paragraph(p, after=3, line=1.0)
    set_run_font(p.add_run(icon), 13.0, bold=True, color=INK)
    set_run_font(p.add_run("  "), 10.5, color=BODY)
    add_hyperlink(p, value, url, size=10.5)
    return p


def add_role_header(cell, company, role, dates, *, before=0, after=7):
    p = cell.add_paragraph()
    set_paragraph(p, before=before, after=after, line=1.0, keep_with_next=True)
    p.paragraph_format.tab_stops.add_tab_stop(
        Twips(DETAIL_WIDTH_DXA - 120), WD_TAB_ALIGNMENT.RIGHT
    )
    set_run_font(p.add_run(company.upper()), 10.8, bold=True, color=INK)
    set_run_font(p.add_run(f"  {role.upper()}"), 9.9, italic=True, color=BODY)
    p.add_run("\t")
    set_run_font(p.add_run(dates), 9.5, color=BODY)
    return p


def add_body(cell, text, *, size=10.0, before=0, after=7, line=1.13, indent=0):
    p = cell.add_paragraph()
    set_paragraph(p, before=before, after=after, line=line)
    p.paragraph_format.left_indent = Inches(indent)
    set_run_font(p.add_run(text), size, color=BODY)
    return p


def add_expertise_line(cell, category, text):
    p = cell.add_paragraph()
    set_paragraph(p, after=4, line=1.10)
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.95)
    fmt.first_line_indent = Inches(-0.95)
    fmt.tab_stops.add_tab_stop(Inches(0.95), WD_TAB_ALIGNMENT.LEFT)
    set_run_font(p.add_run(category.upper()), 8.7, color=ACCENT, font_name=LABEL_FONT)
    p.add_run("\t")
    set_run_font(p.add_run(text), 9.6, color=BODY)
    return p


doc = Document()
configure_styles(doc)
set_document_background(doc)
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = PAGE_MARGIN
section.bottom_margin = PAGE_MARGIN
section.left_margin = PAGE_MARGIN
section.right_margin = PAGE_MARGIN
section.header_distance = Inches(0.15)
section.footer_distance = Inches(0.18)

doc.core_properties.title = "Jacob Malaska - Senior Industrial Designer Resume"
doc.core_properties.author = "Jacob Malaska"
doc.core_properties.subject = "Senior Industrial Designer resume"

bullet_num_id = add_real_bullet_numbering(doc)
table = doc.add_table(rows=5, cols=2)
table.alignment = 0
table.autofit = False

for row in table.rows:
    set_cant_split(row)
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_border(cell)
        set_cell_shading(cell)

# Header row
name_cell, header_cell = table.rows[0].cells
set_cell_margins(name_cell, top=80, bottom=260, start=0, end=180)
set_cell_margins(header_cell, top=80, bottom=260, start=0, end=120)
name_p = clear_paragraph(name_cell.paragraphs[0])
set_paragraph(name_p, line=1.0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
set_run_font(name_p.add_run("JACOB\nMALASKA"), 25.5, color=INK)

title_p = clear_paragraph(header_cell.paragraphs[0])
set_paragraph(title_p, after=7, line=1.0)
set_run_font(title_p.add_run("SENIOR INDUSTRIAL DESIGNER"), 11.7, bold=True, color=INK)
add_contact_line(header_cell, "\u25a3", "redmalaska@gmail.com", "mailto:redmalaska@gmail.com")
add_contact_line(
    header_cell,
    "\u25ce",
    "https://jacob-malaska.github.io/Portfolio-Arcane/",
    "https://jacob-malaska.github.io/Portfolio-Arcane/",
)
add_contact_line(header_cell, "\u260e", "435.216.2040", "tel:+14352162040")

# Experience row
experience_label, experience_cell = table.rows[1].cells
for cell in (experience_label, experience_cell):
    set_cell_border(cell, top={"sz": 6, "color": BORDER})
    set_cell_margins(cell, top=180, bottom=210, start=0, end=120)
set_section_label(experience_label, "Work Experience")
clear_paragraph(experience_cell.paragraphs[0])
experience_cell._element.remove(experience_cell.paragraphs[0]._element)

add_role_header(
    experience_cell,
    "Skullcandy Inc.",
    "Senior Industrial Designer",
    "2019 - PRESENT",
)
for bullet_text in (
    "Lead industrial design across over-ear headphones, true wireless earbuds, gaming "
    "headsets, and accessories for a global consumer audio brand.",
    "Own projects from research and strategy through sketching, CAD, prototyping, "
    "visualization, fit and acoustic validation, DFM, and mass production.",
    "Collaborate with CMF, mechanical, electrical, firmware, acoustics, program "
    "management, and CG teams to ship multiple product generations on lean teams and "
    "fast timelines.",
    "Translate complex technology into approachable product experiences through "
    "tactile controls, ergonomic fit systems, visible structures, and expressive "
    "material stories.",
):
    p = experience_cell.add_paragraph()
    set_paragraph(p, after=5, line=1.13)
    apply_bullet(p, bullet_num_id)
    set_run_font(p.add_run(bullet_text), 10.0, color=BODY)

add_role_header(
    experience_cell,
    "Freelance",
    "Industrial Designer / Consultant",
    "2020 - PRESENT",
    before=9,
)
add_body(
    experience_cell,
    "Design and development across softgoods, jewelry, automotive accessories, and "
    "visualization for clients including Shyft Global, Chums, Kizik, and Ikaika.",
    size=9.9,
    after=9,
    line=1.13,
    indent=0.32,
)
add_role_header(
    experience_cell,
    "Ooblec LLC",
    "Industrial Designer",
    "2017 - 2018",
    before=3,
)
add_body(
    experience_cell,
    "Developed consumer products from concept through prototyping and manufacturer "
    "handoff, spanning inflatables, furniture, and audio equipment.",
    size=9.9,
    after=0,
    line=1.13,
    indent=0.32,
)

# Qualifications row
qualifications_label, qualifications_cell = table.rows[2].cells
for cell in (qualifications_label, qualifications_cell):
    set_cell_border(cell, top={"sz": 6, "color": BORDER})
    set_cell_margins(cell, top=180, bottom=190, start=0, end=120)
set_section_label(qualifications_label, "Qualifications")
clear_paragraph(qualifications_cell.paragraphs[0])
qualifications_cell._element.remove(qualifications_cell.paragraphs[0]._element)
for bullet_text in (
    "9+ years developing consumer products from early research and strategy through mass production.",
    "Experience across consumer electronics, hardgoods, softgoods, jewelry, and audio products.",
    "Deep capability in CAD, visualization, ergonomic validation, DFM, and cross-functional development.",
):
    p = qualifications_cell.add_paragraph()
    set_paragraph(p, after=6, line=1.13)
    apply_bullet(p, bullet_num_id)
    set_run_font(p.add_run(bullet_text), 10.2, color=BODY)

# Education row
education_label, education_cell = table.rows[3].cells
for cell in (education_label, education_cell):
    set_cell_border(cell, top={"sz": 6, "color": BORDER})
    set_cell_margins(cell, top=180, bottom=190, start=0, end=120)
set_section_label(education_label, "Education")
education_p = clear_paragraph(education_cell.paragraphs[0])
set_paragraph(education_p, after=5, line=1.0, keep_with_next=True)
education_p.paragraph_format.tab_stops.add_tab_stop(
    Twips(DETAIL_WIDTH_DXA - 120), WD_TAB_ALIGNMENT.RIGHT
)
set_run_font(education_p.add_run("BFA, INDUSTRIAL DESIGN"), 10.8, bold=True, color=INK)
education_p.add_run("\t")
set_run_font(education_p.add_run("2015 - 2019"), 9.5, color=BODY)
school_p = education_cell.add_paragraph()
set_paragraph(school_p, after=7, line=1.0, keep_with_next=True)
set_run_font(school_p.add_run("Brigham Young University"), 10.0, color=BODY)
add_body(
    education_cell,
    "Sponsored projects for Black Diamond and Native Shoes. Multidisciplinary and "
    "entrepreneurial training through the Crocker Innovation Fellowship.",
    size=9.8,
    after=0,
    line=1.13,
)

# Expertise row
expertise_label, expertise_cell = table.rows[4].cells
for cell in (expertise_label, expertise_cell):
    set_cell_border(cell, top={"sz": 6, "color": BORDER})
    set_cell_margins(cell, top=180, bottom=0, start=0, end=120)
set_section_label(expertise_label, "Expertise")
clear_paragraph(expertise_cell.paragraphs[0])
expertise_cell._element.remove(expertise_cell.paragraphs[0]._element)
for category, detail in (
    ("Design", "Research and methodology / Sketching / Hardgoods / Softgoods development"),
    ("Production", "DFM / Prototyping / Pattern making / Woodworking"),
    ("Tools", "SolidWorks / Blender / KeyShot / Adobe Creative Suite"),
    ("Media", "3D visualization / Animation / Photography / Video editing / Audio and sound design"),
):
    add_expertise_line(expertise_cell, category, detail)

# Enforce fixed, matching OOXML geometry after all rows and cells exist.
apply_table_geometry(
    table,
    [LABEL_WIDTH_DXA, DETAIL_WIDTH_DXA],
    table_width_dxa=CONTENT_WIDTH_DXA,
    indent_dxa=0,
    cell_margins_dxa={"top": 120, "bottom": 120, "start": 0, "end": 120},
)

# Restore deliberate per-section vertical padding after the geometry pass.
for row_idx, row in enumerate(table.rows):
    if row_idx == 0:
        margins = dict(top=80, bottom=260, start=0, end=120)
        table.rows[row_idx].cells[0] and set_cell_margins(
            table.rows[row_idx].cells[0], top=80, bottom=260, start=0, end=180
        )
        set_cell_margins(table.rows[row_idx].cells[1], **margins)
        continue
    elif row_idx == 4:
        margins = dict(top=180, bottom=0, start=0, end=120)
    elif row_idx == 1:
        margins = dict(top=180, bottom=210, start=0, end=120)
    else:
        margins = dict(top=180, bottom=190, start=0, end=120)
    for cell in row.cells:
        set_cell_margins(cell, **margins)
        set_cell_shading(cell)

# Quiet one-page update footer; no decorative rule or running chrome.
footer_p = section.footer.paragraphs[0]
clear_paragraph(footer_p)
set_paragraph(footer_p, line=1.0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
set_run_font(footer_p.add_run("UPDATED 07.2026"), 8.0, color=MUTED, font_name=LABEL_FONT)

doc.save(RAW_OUT)
print(RAW_OUT)
print(FINAL_OUT)
